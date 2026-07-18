from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output"
OUT_DIR.mkdir(exist_ok=True)
DOCX_PATH = OUT_DIR / "采摘机器人技术报告_雏形.docx"
PCB_IMAGE = Path(r"C:\Users\ZhouJL\AppData\Local\Temp\codex-clipboard-eb05309e-bd41-421f-97d5-62bf8c468f57.png")

ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
LIGHT_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_run_font(run, name="Microsoft YaHei", size=None, color=None, bold=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=10, color=color, bold=bold)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = p.add_run(text)
    set_run_font(run, size=16 if level == 1 else 13, color=ACCENT if level < 3 else DARK, bold=True)
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=10.5, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, size=10.5)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        run = p.add_run(item)
        set_run_font(run, size=10.5)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_widths(table, [6.35])
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, color=DARK, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_run_font(r2, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], LIGHT_FILL)
        set_cell_text(hdr[i], h, bold=True, color=DARK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, ACCENT),
        ("Heading 3", 12, DARK),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("技术报告初稿")
    set_run_font(r, size=12, color=MUTED, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("2026 采摘机器人项目技术报告")
    set_run_font(r, size=25, color=RGBColor(0, 0, 0), bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    r = subtitle.add_run("基于 STM32F407VET6 与 K230 视觉模块的自主采摘、识别与运输系统")
    set_run_font(r, size=13, color=MUTED)

    rows = [
        ("项目阶段", "电控与视觉原型已具备，底盘/机械整机联调待完成"),
        ("报告版本", "雏形版，供 2026-07-10 前继续补充"),
        ("主控平台", "STM32F407VET6，HAL + FreeRTOS"),
        ("视觉平台", "K230，YOLOv8 / fruit7.kmodel"),
        ("资料来源", "当前工程源码、比赛规则 PDF、主控扩展板 PCB 截图"),
    ]
    add_table(doc, ["项目", "内容"], rows, [1.35, 5.0])
    add_callout(
        doc,
        "当前写作原则",
        "本报告先形成可提交的技术框架：已完成部分按当前代码和硬件进度写实，导航闭环、底盘整车测试和机械结构章节明确标为待完善，避免把尚未验证的内容写成最终结果。",
    )


def build_doc():
    doc = Document()
    configure_doc(doc)
    add_cover(doc)

    add_heading(doc, "目录", 1)
    add_bullets(doc, [
        "1 项目背景与比赛需求",
        "2 系统总体方案",
        "3 电控系统设计",
        "4 K230 视觉识别与串口协同",
        "5 主控扩展板硬件设计",
        "6 机械结构设计（待补充）",
        "7 当前进度、问题与后续计划",
        "8 附录：工程证据与术语说明",
    ])

    add_heading(doc, "1 项目背景与比赛需求", 1)
    add_para(
        doc,
        "本项目面向 2026 采摘机器人赛项，目标是在标准化场地中完成果蔬识别、采摘、坏果处理、运输和返回等自主作业。根据比赛规则，机器人需要在 5200 mm x 3000 mm 场地内全程无遥控运行，并在 10 分钟内尽可能完成各区域任务。"
    )
    add_table(
        doc,
        ["规则点", "对本项目的设计影响"],
        [
            ("全程自主作业", "控制系统必须具备自主导航、状态判断和任务调度能力，不能依赖遥控操作。"),
            ("目标识别与坏果处理", "视觉模块需要区分成熟、未成熟和坏果，并将结果传给 STM32 执行抓取或丢弃策略。"),
            ("语音播报", "系统需要在识别、作业状态或库存播报时触发语音模块输出。"),
            ("500 mm x 500 mm 垂直投影限制", "底盘、扩展板、机械臂和果蔬存放装置需要控制整体尺寸。"),
            ("禁止麦克纳姆轮/全向轮/履带式底盘", "移动机构需采用更接近农业环境的普通轮式底盘方案。"),
        ],
        [1.65, 4.7],
    )

    add_heading(doc, "2 系统总体方案", 1)
    add_para(
        doc,
        "系统采用“STM32 主控 + K230 视觉 + 多传感器定位 + 步进电机底盘 + 舵机执行机构”的分层结构。STM32F407VET6 负责底盘控制、传感器采集、里程计更新、舵机动作和语音播报；K230 负责摄像头图像采集、YOLO 目标检测、二维码读取和目标状态判断。"
    )
    add_table(
        doc,
        ["层级", "主要模块", "当前状态"],
        [
            ("感知层", "HWT101 陀螺仪、TOF200F 测距、K230 摄像头/YOLO", "陀螺仪、TOF、视觉识别代码已具备；整车场地测试待做。"),
            ("决策层", "导航坐标更新、任务状态、二维码顺序、坏果优先处理", "导航坐标框架已写入 STM32；完整路径规划算法待补充。"),
            ("执行层", "Emm V5 步进电机、舵机、语音板、OLED 显示", "步进控制、里程计、舵机、语音模块均有驱动代码。"),
            ("硬件层", "STM32 主控扩展板、串口/I2C/PWM 接口、电源接口", "扩展板已完成绘制和打板，需继续做上电与接口联调记录。"),
        ],
        [1.2, 2.55, 2.6],
    )
    add_callout(
        doc,
        "工作流程雏形",
        "起始区初始化 -> HWT101 建立 yaw 零点 -> 里程计周期读取位移 -> 根据地图坐标点移动到目标区域 -> TOF 辅助距离判断 -> K230 识别目标成熟度/坏果 -> STM32 控制舵机抓取或丢弃 -> 完成运输并返回。",
    )

    add_heading(doc, "3 电控系统设计", 1)
    add_heading(doc, "3.1 主控与接口分配", 2)
    add_table(
        doc,
        ["功能", "接口/外设", "工程证据", "说明"],
        [
            ("语音播报", "USART3", "Core/Src/voice.c", "JQ8x00 语音模块按编号播放水果、成熟度、坏果和数字音频。"),
            ("陀螺仪姿态", "UART4，115200", "Core/Src/hwt101_hal.c", "解析 HWT101 0x55 0x53 角度包，得到 roll/pitch/yaw。"),
            ("步进电机控制", "USART2，921600", "Core/Src/bujin.c", "封装 Emm V5 使能、速度、位置、同步运动和读参数指令。"),
            ("里程计", "USART2 回包", "Core/Src/odometer.c", "读取 S_CPOS，按 360/65536 换算角度并积分位移。"),
            ("TOF 红外测距", "USART1，115200", "Core/Src/tof200f.c", "发送单次测距命令，回包距离值按 /10 得到 cm。"),
            ("舵机抓取", "TIM3_CH4 / TIM8_CH1", "Core/Src/servo.c", "支持 270 度和 180 度舵机角度控制。"),
            ("OLED/调试显示", "I2C2", "Core/Src/oled.c", "用于显示坐标、姿态或调试信息。"),
            ("蓝牙调试", "USART6", "Core/Src/bluetooth.c", "作为调试或临时通信接口。"),
        ],
        [1.25, 1.15, 1.45, 2.5],
    )

    add_heading(doc, "3.2 软件任务与数据流", 2)
    add_para(
        doc,
        "STM32 工程基于 HAL 与 CMSIS-RTOS2/FreeRTOS。main.c 中完成 GPIO、UART、TIM、I2C 初始化，随后初始化 HWT101、蓝牙、导航零点、里程计和 OLED，并启动 RTOS 调度。freertos.c 中已创建多个任务，其中里程计任务每 20 ms 调用 Odometer_Update()，持续读取电机当前位置用于里程估计。"
    )
    add_bullets(doc, [
        "HWT101 数据由 UART4 中断解析，更新 yaw 角；导航模块以启动时 yaw 为零点。",
        "里程计从步进驱动器读取当前位置，计算角度增量和线位移，再调用 Navigation_UpdateByDelta() 更新坐标。",
        "步进电机控制封装了互斥量保护，避免 RTOS 多任务同时通过 USART2 发送指令时发生冲突。",
        "当前 Keil 构建日志显示 Code=20104，0 Error(s)，0 Warning(s)，说明 STM32 侧代码可编译通过。",
    ])

    add_heading(doc, "3.3 导航方案现状", 2)
    add_para(
        doc,
        "当前导航采用惯性导航与里程计融合的雏形：HWT101 提供车体航向角，里程计提供沿车头方向的增量位移，navigation.c 中维护 g_nav_x_cm、g_nav_y_cm、g_nav_yaw_deg 三个全局估计量。地图宽高在 navigation.h 中定义为 520 cm x 240 cm，起点中心暂定为 (25 cm, 25 cm)。"
    )
    add_callout(
        doc,
        "未验证点",
        "底盘尚未完成，因此点对点导航、转向 PID 参数、地图坐标误差累积、TOF 避障策略还没有整车实测。后续需要在底盘完成后先做直线、原地转向、定点停车三类标定，再考虑完整路径规划。",
    )

    add_heading(doc, "4 K230 视觉识别与串口协同", 1)
    add_para(
        doc,
        "K230 侧使用 YOLOv8 模型 fruit7.kmodel 进行目标检测，类别覆盖番茄、辣椒、南瓜、洋葱、苹果、梨以及坏果，并区分成熟/未成熟状态。程序中通过连续识别计数降低误检，确认目标后通过串口向下位机发送 voice、arm、QR 等命令。"
    )
    add_table(
        doc,
        ["能力", "实现方式", "当前状态"],
        [
            ("成熟/未成熟/坏果识别", "YOLOv8 检测类别 + transform_list 映射语音码和状态码", "代码已具备，模型现场置信度仍需继续采样验证。"),
            ("二维码任务顺序", "解析二维码水果名称和位置编号，生成 fruit_list 与 number", "已具备解析和串口发送 QR 编号逻辑。"),
            ("坏果优先处理", "检测到 bad 类别时优先选择坏果框，发送抓取/丢弃指令", "已在视觉逻辑中体现。"),
            ("与 STM32 通信", "K230 串口发送 voice:x、arm:x、QR:x 等字符串", "协议框架已具备，整机串口联调仍需补记录。"),
            ("调试状态", "当前 main.py 中 FORCE_TASK = \"send\"", "比赛联调前应改为 None，使 K230 从串口读取 scan/send/pour 任务。"),
        ],
        [1.55, 2.65, 2.15],
    )

    add_heading(doc, "5 主控扩展板硬件设计", 1)
    add_para(
        doc,
        "硬件部分已完成 STM32 主控扩展板绘制和打板。扩展板围绕主控最小系统、电源输入、舵机接口、K230 通信接口、蓝牙模块、OLED I2C、TOF、语音模块和多路电源输出进行接口化设计，便于后续底盘、视觉和执行机构集中接线。"
    )
    if PCB_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(PCB_IMAGE), width=Inches(6.35))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run("图 1 主控扩展板 PCB 初版截图")
        set_run_font(r, size=9.5, color=MUTED, bold=True)
    else:
        add_callout(doc, "图片缺失", "未找到 PCB 截图文件，后续可补充主控扩展板正反面照片或 KiCad/AD 渲染图。")

    add_table(
        doc,
        ["硬件子项", "完成情况", "后续需要补充的验证"],
        [
            ("主控扩展板原理/PCB", "已完成绘制和打板", "上电电压检查、短路检查、关键接口电平测试。"),
            ("电源接口", "板上预留多路 VCC/GND/12V 等接口", "确认舵机、电机、K230 和 STM32 共地及电流裕量。"),
            ("通信接口", "预留 K230、蓝牙、语音、TOF、OLED 等接口", "逐个接口做收发测试并记录波特率/线序。"),
            ("机械安装", "板上有安装孔和外形边界", "结合底盘结构确认固定孔位与线束走向。"),
        ],
        [1.55, 2.15, 2.65],
    )

    add_heading(doc, "6 机械结构设计（待补充）", 1)
    add_callout(
        doc,
        "本章暂留空",
        "机械部分目前不强行补写。后续建议补充底盘结构、轮距/轴距、抓取机构自由度、果蔬存放装置、坏果丢弃机构、尺寸校核和关键零件图。"
    )
    add_table(
        doc,
        ["待补内容", "建议写法"],
        [
            ("底盘方案", "说明普通轮式底盘、驱动轮数量、轮距、转向方式和为什么符合规则。"),
            ("抓取机构", "说明舵机数量、夹爪/拨杆结构、与 K230 视觉中心对准的动作流程。"),
            ("运输/存放结构", "说明果蔬进入收集区后的暂存方式，以及坏果丢弃路径。"),
            ("尺寸与可靠性", "给出整机垂直投影是否小于 500 mm x 500 mm，说明重心和防碰撞设计。"),
        ],
        [1.55, 4.8],
    )

    add_heading(doc, "7 当前进度、问题与后续计划", 1)
    add_table(
        doc,
        ["模块", "已完成", "待完成/风险"],
        [
            ("语音播报", "USART3 驱动 JQ8x00，语音编号 1-20/99 已封装", "需要补现场播报视频或测试记录。"),
            ("姿态感知", "HWT101 yaw/roll/pitch 解析与初始化完成", "需要标定安装方向和 yaw 漂移。"),
            ("步进与里程计", "Emm V5 控制、同步运动、S_CPOS 读取、里程积分已完成", "底盘未出，轮径/滑移/编码误差未实测。"),
            ("TOF 测距", "TOF200F 串口解析与 cm 换算完成", "避障/定位触发逻辑还需和路径规划结合。"),
            ("舵机抓取", "270 度、180 度舵机 PWM 控制完成", "需和机械夹爪、K230 中心对准做联调。"),
            ("K230 视觉", "YOLO 识别、二维码解析、串口协议、坏果优先逻辑已具备", "联调前取消 FORCE_TASK 调试开关，补充板端测试截图。"),
            ("硬件 PCB", "主控扩展板已绘制和打板", "上电、接口、负载能力和抗干扰测试待记录。"),
            ("机械结构", "暂未写入", "需要机械同学补结构图、尺寸、机构动作和装配说明。"),
        ],
        [1.25, 2.55, 2.55],
    )
    add_heading(doc, "7.1 周五前最低可交付补充清单", 2)
    add_bullets(doc, [
        "补 1 张整机或底盘设计草图，即使是手绘也比空白强。",
        "补 1 张 K230 识别成熟/未成熟/坏果的运行截图。",
        "补 1 段语音播报和 1 段舵机抓取动作测试照片/视频截图。",
        "补 PCB 实物照片和一次万用表上电测试表。",
        "把导航章节改成“方案与待验证”，不要写成已经完成整场导航。",
    ])

    add_heading(doc, "8 附录：工程证据与术语说明", 1)
    add_table(
        doc,
        ["证据", "位置/说明"],
        [
            ("STM32 工程", "D:/codexproject/智慧农业赛道/厂里大运/stm32f407vet6"),
            ("Keil 构建", "MDK-ARM/build.log：0 Error(s)，0 Warning(s)"),
            ("主函数初始化", "Core/Src/main.c：UART4、UART5、USART1/2/3/6、TIM3/7/8、I2C2 初始化"),
            ("导航模块", "Core/Src/navigation.c 与 Core/Inc/navigation.h"),
            ("视觉模块", "D:/codexproject/智慧农业赛道/厂里大运/视觉代码/main.py"),
            ("比赛规则", "2026采摘机器人比赛规则.pdf：场地、任务、尺寸与自主作业要求"),
        ],
        [1.55, 4.8],
    )
    add_heading(doc, "重点英文/缩写", 2)
    add_table(
        doc,
        ["词汇", "含义"],
        [
            ("TOF", "Time of Flight，飞行时间测距，用于获取距离。"),
            ("YOLO", "You Only Look Once，一类实时目标检测算法。"),
            ("UART / USART", "串口通信接口，用于模块之间传输数据。"),
            ("Odometer", "里程计，用于估计轮子或车体运动距离。"),
            ("PWM", "Pulse Width Modulation，脉宽调制，用于控制舵机角度。"),
            ("PID", "比例-积分-微分控制，用于姿态或速度闭环调节。"),
        ],
        [1.35, 5.0],
    )

    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build_doc()
