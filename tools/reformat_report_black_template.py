from pathlib import Path
from zipfile import ZipFile
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "output"
OUT_DIR.mkdir(exist_ok=True)
OUT_PATH = OUT_DIR / "采摘机器人技术报告_模板重排版.docx"
PCB_IMAGE = Path(r"C:\Users\ZhouJL\AppData\Local\Temp\codex-clipboard-eb05309e-bd41-421f-97d5-62bf8c468f57.png")

BLACK = RGBColor(0, 0, 0)
GRAY_FILL = "F2F2F2"


def set_run_font(run, east="宋体", ascii_font="Times New Roman", size=None, bold=None):
    run.font.name = ascii_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = BLACK


def set_paragraph(p, align=None, before=None, after=None, line_spacing=None, first_indent=True):
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    if before is not None:
        fmt.space_before = Pt(before)
    if after is not None:
        fmt.space_after = Pt(after)
    if line_spacing is not None:
        fmt.line_spacing = line_spacing
    if first_indent:
        fmt.first_line_indent = Cm(0.35)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill=GRAY_FILL):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_widths(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.width = Inches(width)
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def add_text(p, text, east="宋体", ascii_font="Times New Roman", size=12, bold=False):
    run = p.add_run(text)
    set_run_font(run, east=east, ascii_font=ascii_font, size=size, bold=bold)
    return run


def add_para(doc, text, first_indent=True):
    p = doc.add_paragraph()
    set_paragraph(p, after=0, line_spacing=1.3, first_indent=first_indent)
    add_text(p, text, size=12)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    set_paragraph(p, before=6 if level == 1 else 3, after=3, line_spacing=1.3, first_indent=False)
    if level == 1:
        add_text(p, text, east="黑体", size=14, bold=False)
    elif level == 2:
        add_text(p, text, east="黑体", size=12, bold=True)
    else:
        add_text(p, text, east="宋体", size=12, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=3, line_spacing=1.3, first_indent=False)
    add_text(p, text, east="黑体", size=10.5)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell)
        p = cell.paragraphs[0]
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, line_spacing=1.15, first_indent=False)
        add_text(p, header, east="黑体", size=10.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            set_paragraph(p, after=0, line_spacing=1.15, first_indent=False)
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_text(p, str(value), size=10.5)
    doc.add_paragraph()
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.first_line_indent = Cm(0.35)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[style_name]
        style.font.color.rgb = BLACK
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if style_name != "Heading 3" else "宋体")
        style.font.size = Pt(14 if style_name == "Heading 1" else 12)
        style.paragraph_format.line_spacing = 1.3
        style.paragraph_format.first_line_indent = Cm(0)


def add_cover(doc):
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, before=18, after=8, first_indent=False)
    add_text(p, "第二十八届中国机器人及人工智能大赛", east="宋体", size=26, bold=True)

    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT, after=22, first_indent=False)
    add_text(p, "——采摘机器人设计制作竞赛", east="黑体", size=18)

    for ch in "技术报告":
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, first_indent=False)
        add_text(p, ch, east="宋体", size=36, bold=True)

    doc.add_paragraph()
    for line in [
        "团队名称：__________",
        "团队编号：__________",
        "作品名称：采摘机器人",
        "作品编号：__________",
    ]:
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, after=3, line_spacing=1.5, first_indent=False)
        p.paragraph_format.left_indent = Cm(3.0)
        add_text(p, line, east="仿宋", size=18)
    doc.add_page_break()


def add_abstract(doc):
    p = doc.add_paragraph()
    set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, first_indent=False)
    add_text(p, "摘  要", east="黑体", size=14, bold=True)

    add_para(
        doc,
        "本项目面向2026采摘机器人赛项，围绕果蔬自主识别、采摘、坏果处理、运输和返回等任务，设计了一套基于STM32F407VET6主控与K230视觉模块的采摘机器人控制系统。系统以STM32F407VET6为下位机，完成传感器采集、步进电机控制、舵机执行、语音播报、里程计更新与导航坐标维护；以K230为上位视觉模块，部署YOLOv8目标检测模型，用于区分成熟水果、未成熟水果和坏水果，并通过串口与主控协同完成精准抓取流程。"
    )
    add_para(
        doc,
        "目前电控部分已完成语音板播报、HWT101陀螺仪数据获取、步进电机控制、里程计数据读取、TOF红外测距、舵机控制及K230串口协同框架；硬件部分已完成主控扩展板绘制和打板。由于底盘和机械整机尚未完成，惯性导航、点对点路径规划、机械抓取结构和整车任务闭环仍处于待验证阶段。本文档先形成项目技术报告雏形，明确已完成内容、系统方案、当前风险与后续补充方向。"
    )
    p = doc.add_paragraph()
    set_paragraph(p, after=12, line_spacing=1.3, first_indent=False)
    add_text(p, "关键词：", east="黑体", size=12, bold=True)
    add_text(p, "采摘机器人；STM32F407VET6；K230视觉模块；YOLOv8；惯性导航；里程计", size=12)
    doc.add_page_break()


def add_directory(doc):
    add_heading(doc, "目录", 1)
    for line in [
        "一、引言",
        "二、系统总体方案",
        "三、电控系统设计",
        "四、K230视觉识别与串口协同",
        "五、主控扩展板硬件设计",
        "六、机械结构设计（待补充）",
        "七、当前进度、问题与后续计划",
        "八、附录：工程证据与术语说明",
    ]:
        p = doc.add_paragraph()
        set_paragraph(p, after=2, line_spacing=1.3, first_indent=False)
        add_text(p, line, size=12)
    doc.add_page_break()


def build_report():
    doc = Document()
    configure_doc(doc)
    add_cover(doc)
    add_abstract(doc)
    add_directory(doc)

    add_heading(doc, "一、引言", 1)
    add_para(
        doc,
        "本项目面向2026采摘机器人赛项，目标是在标准化场地中完成果蔬识别、采摘、坏果处理、运输和返回等自主作业。根据比赛规则，机器人需要在5200mm×3000mm场地内全程无遥控运行，并在10分钟内尽可能完成各区域任务。"
    )
    add_table(
        doc,
        ["规则点", "对本项目的设计影响"],
        [
            ("全程自主作业", "控制系统必须具备自主导航、状态判断和任务调度能力，不能依赖遥控操作。"),
            ("目标识别与坏果处理", "视觉模块需要区分成熟、未成熟和坏果，并将结果传给STM32执行抓取或丢弃策略。"),
            ("语音播报", "系统需要在识别、作业状态或库存播报时触发语音模块输出。"),
            ("尺寸限制", "底盘、扩展板、机械臂和果蔬存放装置需要控制整体尺寸。"),
            ("轮式底盘约束", "移动机构需采用更接近农业环境的普通轮式底盘方案。"),
        ],
        [1.55, 4.75],
    )

    add_heading(doc, "二、系统总体方案", 1)
    add_para(
        doc,
        "系统采用“STM32主控 + K230视觉 + 多传感器定位 + 步进电机底盘 + 舵机执行机构”的分层结构。STM32F407VET6负责底盘控制、传感器采集、里程计更新、舵机动作和语音播报；K230负责摄像头图像采集、YOLO目标检测、二维码读取和目标状态判断。"
    )
    add_table(
        doc,
        ["层级", "主要模块", "当前状态"],
        [
            ("感知层", "HWT101陀螺仪、TOF200F测距、K230摄像头/YOLO", "陀螺仪、TOF、视觉识别代码已具备；整车场地测试待做。"),
            ("决策层", "导航坐标更新、任务状态、二维码顺序、坏果优先处理", "导航坐标框架已写入STM32；完整路径规划算法待补充。"),
            ("执行层", "Emm V5步进电机、舵机、语音板、OLED显示", "步进控制、里程计、舵机、语音模块均有驱动代码。"),
            ("硬件层", "STM32主控扩展板、串口/I2C/PWM接口、电源接口", "扩展板已完成绘制和打板，需继续做上电与接口联调记录。"),
        ],
        [1.15, 2.55, 2.6],
    )
    add_para(
        doc,
        "工作流程初步设计为：起始区初始化，HWT101建立yaw零点，里程计周期读取位移，根据地图坐标点移动到目标区域，TOF辅助距离判断，K230识别目标成熟度或坏果，STM32控制舵机抓取或丢弃，最后完成运输并返回。"
    )

    add_heading(doc, "三、电控系统设计", 1)
    add_heading(doc, "3.1 主控与接口分配", 2)
    add_table(
        doc,
        ["功能", "接口/外设", "工程证据", "说明"],
        [
            ("语音播报", "USART3", "Core/Src/voice.c", "JQ8x00语音模块按编号播放水果、成熟度、坏果和数字音频。"),
            ("陀螺仪姿态", "UART4，115200", "Core/Src/hwt101_hal.c", "解析HWT101角度包，得到roll、pitch、yaw。"),
            ("步进电机控制", "USART2，921600", "Core/Src/bujin.c", "封装Emm V5使能、速度、位置、同步运动和读参数指令。"),
            ("里程计", "USART2回包", "Core/Src/odometer.c", "读取S_CPOS，按360/65536换算角度并积分位移。"),
            ("TOF红外测距", "USART1，115200", "Core/Src/tof200f.c", "发送单次测距命令，回包距离值按/10得到cm。"),
            ("舵机抓取", "TIM3_CH4 / TIM8_CH1", "Core/Src/servo.c", "支持270度和180度舵机角度控制。"),
            ("OLED显示", "I2C2", "Core/Src/oled.c", "用于显示坐标、姿态或调试信息。"),
        ],
        [1.15, 1.05, 1.5, 2.6],
    )
    add_heading(doc, "3.2 软件任务与数据流", 2)
    add_para(
        doc,
        "STM32工程基于HAL与CMSIS-RTOS2/FreeRTOS。main.c中完成GPIO、UART、TIM、I2C初始化，随后初始化HWT101、蓝牙、导航零点、里程计和OLED，并启动RTOS调度。freertos.c中已创建多个任务，其中里程计任务每20ms调用Odometer_Update()，持续读取电机当前位置用于里程估计。"
    )
    add_para(
        doc,
        "HWT101数据由UART4中断解析并更新yaw角；里程计从步进驱动器读取当前位置，计算角度增量和线位移，再调用Navigation_UpdateByDelta()更新坐标。步进电机控制代码中加入互斥量保护，避免多任务同时通过USART2发送指令造成冲突。当前Keil构建日志显示0 Error(s)，0 Warning(s)，说明STM32侧代码能够编译通过。"
    )
    add_heading(doc, "3.3 导航方案现状", 2)
    add_para(
        doc,
        "当前导航采用惯性导航与里程计融合的雏形：HWT101提供车体航向角，里程计提供沿车头方向的增量位移，navigation.c中维护g_nav_x_cm、g_nav_y_cm、g_nav_yaw_deg三个全局估计量。地图宽高在navigation.h中定义为520cm×240cm，起点中心暂定为(25cm, 25cm)。"
    )
    add_para(
        doc,
        "由于底盘尚未完成，点对点导航、转向PID参数、地图坐标误差累积、TOF避障策略还没有整车实测。后续需要在底盘完成后先做直线、原地转向、定点停车三类标定，再考虑完整路径规划。"
    )

    add_heading(doc, "四、K230视觉识别与串口协同", 1)
    add_para(
        doc,
        "K230侧使用YOLOv8模型fruit7.kmodel进行目标检测，类别覆盖番茄、辣椒、南瓜、洋葱、苹果、梨以及坏果，并区分成熟和未成熟状态。程序中通过连续识别计数降低误检，确认目标后通过串口向下位机发送voice、arm、QR等命令。"
    )
    add_table(
        doc,
        ["能力", "实现方式", "当前状态"],
        [
            ("成熟/未成熟/坏果识别", "YOLOv8检测类别和transform_list映射语音码、状态码", "代码已具备，模型现场置信度仍需继续采样验证。"),
            ("二维码任务顺序", "解析二维码水果名称和位置编号，生成fruit_list与number", "已具备解析和串口发送QR编号逻辑。"),
            ("坏果优先处理", "检测到bad类别时优先选择坏果框，发送抓取/丢弃指令", "已在视觉逻辑中体现。"),
            ("与STM32通信", "K230串口发送voice:x、arm:x、QR:x等字符串", "协议框架已具备，整机串口联调仍需补记录。"),
            ("调试状态", "当前main.py中FORCE_TASK = \"send\"", "比赛联调前应改为None，使K230从串口读取scan/send/pour任务。"),
        ],
        [1.55, 2.55, 2.2],
    )

    add_heading(doc, "五、主控扩展板硬件设计", 1)
    add_para(
        doc,
        "硬件部分已完成STM32主控扩展板绘制和打板。扩展板围绕主控最小系统、电源输入、舵机接口、K230通信接口、蓝牙模块、OLED I2C、TOF、语音模块和多路电源输出进行接口化设计，便于后续底盘、视觉和执行机构集中接线。"
    )
    if PCB_IMAGE.exists():
        p = doc.add_paragraph()
        set_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False)
        p.add_run().add_picture(str(PCB_IMAGE), width=Inches(5.8))
        add_caption(doc, "图5.1 主控扩展板PCB初版截图")
    add_table(
        doc,
        ["硬件子项", "完成情况", "后续需要补充的验证"],
        [
            ("主控扩展板原理/PCB", "已完成绘制和打板", "上电电压检查、短路检查、关键接口电平测试。"),
            ("电源接口", "板上预留多路VCC/GND/12V等接口", "确认舵机、电机、K230和STM32共地及电流裕量。"),
            ("通信接口", "预留K230、蓝牙、语音、TOF、OLED等接口", "逐个接口做收发测试并记录波特率和线序。"),
            ("机械安装", "板上有安装孔和外形边界", "结合底盘结构确认固定孔位与线束走向。"),
        ],
        [1.5, 2.15, 2.65],
    )

    add_heading(doc, "六、机械结构设计（待补充）", 1)
    add_para(
        doc,
        "机械部分目前先保留为待补充章节。后续建议补充底盘结构、轮距/轴距、抓取机构自由度、果蔬存放装置、坏果丢弃机构、尺寸校核和关键零件图。"
    )
    add_table(
        doc,
        ["待补内容", "建议写法"],
        [
            ("底盘方案", "说明普通轮式底盘、驱动轮数量、轮距、转向方式和为什么符合规则。"),
            ("抓取机构", "说明舵机数量、夹爪/拨杆结构、与K230视觉中心对准的动作流程。"),
            ("运输/存放结构", "说明果蔬进入收集区后的暂存方式，以及坏果丢弃路径。"),
            ("尺寸与可靠性", "给出整机垂直投影是否小于500mm×500mm，说明重心和防碰撞设计。"),
        ],
        [1.5, 4.8],
    )

    add_heading(doc, "七、当前进度、问题与后续计划", 1)
    add_table(
        doc,
        ["模块", "已完成", "待完成/风险"],
        [
            ("语音播报", "USART3驱动JQ8x00，语音编号1-20/99已封装", "需要补现场播报视频或测试记录。"),
            ("姿态感知", "HWT101 yaw/roll/pitch解析与初始化完成", "需要标定安装方向和yaw漂移。"),
            ("步进与里程计", "Emm V5控制、同步运动、S_CPOS读取、里程积分已完成", "底盘未出，轮径、滑移、编码误差未实测。"),
            ("TOF测距", "TOF200F串口解析与cm换算完成", "避障/定位触发逻辑还需和路径规划结合。"),
            ("舵机抓取", "270度、180度舵机PWM控制完成", "需和机械夹爪、K230中心对准做联调。"),
            ("K230视觉", "YOLO识别、二维码解析、串口协议、坏果优先逻辑已具备", "联调前取消FORCE_TASK调试开关，补充板端测试截图。"),
            ("硬件PCB", "主控扩展板已绘制和打板", "上电、接口、负载能力和抗干扰测试待记录。"),
            ("机械结构", "暂未写入", "需要机械同学补结构图、尺寸、机构动作和装配说明。"),
        ],
        [1.15, 2.55, 2.6],
    )
    add_heading(doc, "7.1 周五前最低可交付补充清单", 2)
    for item in [
        "补1张整机或底盘设计草图，即使是手绘也比空白强。",
        "补1张K230识别成熟/未成熟/坏果的运行截图。",
        "补1段语音播报和1段舵机抓取动作测试照片或视频截图。",
        "补PCB实物照片和一次万用表上电测试表。",
        "把导航章节保持为“方案与待验证”，不要写成已经完成整场导航。",
    ]:
        add_para(doc, item)

    add_heading(doc, "八、附录：工程证据与术语说明", 1)
    add_table(
        doc,
        ["证据", "位置/说明"],
        [
            ("STM32工程", "D:/codexproject/智慧农业赛道/厂里大运/stm32f407vet6"),
            ("Keil构建", "MDK-ARM/build.log：0 Error(s)，0 Warning(s)"),
            ("主函数初始化", "Core/Src/main.c：UART4、UART5、USART1/2/3/6、TIM3/7/8、I2C2初始化"),
            ("导航模块", "Core/Src/navigation.c与Core/Inc/navigation.h"),
            ("视觉模块", "D:/codexproject/智慧农业赛道/厂里大运/视觉代码/main.py"),
            ("比赛规则", "2026采摘机器人比赛规则.pdf：场地、任务、尺寸与自主作业要求"),
        ],
        [1.45, 4.85],
    )
    add_heading(doc, "8.1 重点英文与缩写", 2)
    add_table(
        doc,
        ["词汇", "含义"],
        [
            ("TOF", "Time of Flight，飞行时间测距，用于获取距离。"),
            ("YOLO", "You Only Look Once，一类实时目标检测算法。"),
            ("UART/USART", "串口通信接口，用于模块之间传输数据。"),
            ("Odometer", "里程计，用于估计轮子或车体运动距离。"),
            ("PWM", "Pulse Width Modulation，脉宽调制，用于控制舵机角度。"),
            ("PID", "比例-积分-微分控制，用于姿态或速度闭环调节。"),
        ],
        [1.35, 4.95],
    )

    doc.save(OUT_PATH)
    strip_non_black_text_colors(OUT_PATH)
    print(OUT_PATH)


def strip_non_black_text_colors(path):
    # Keep table fills, but force all explicit text color declarations to black.
    import tempfile
    import shutil

    tmp = Path(tempfile.mkdtemp())
    try:
        with ZipFile(path, "r") as zin:
            zin.extractall(tmp)
        for rel in ("word/document.xml", "word/styles.xml"):
            p = tmp / rel
            if p.exists():
                text = p.read_text(encoding="utf-8", errors="ignore")
                text = re.sub(r'w:color w:val="[0-9A-Fa-f]+"', 'w:color w:val="000000"', text)
                p.write_text(text, encoding="utf-8")
        backup = path.with_suffix(".tmp.docx")
        with ZipFile(backup, "w") as zout:
            for file in tmp.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp).as_posix())
        backup.replace(path)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


if __name__ == "__main__":
    build_report()
